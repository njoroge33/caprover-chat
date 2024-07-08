import solara as sol
from typing import List, TypedDict
from docx import Document
import time
import threading
from pathlib import Path
from langchain.llms import Ollama

class MessageDict(TypedDict):
    role: str
    content: str

# Initialize reactive messages list with the first assistant message
initial_message = [{"role": "assistant", "content": "What do you want to search for on Wikipedia?"}]
messages: sol.Reactive[List[MessageDict]] = sol.reactive(initial_message)
typing_complete: sol.Reactive[bool] = sol.reactive(False)

ollama_client = Ollama(base_url="http://ollama.captain.localhost/", model="llama3:8b")

# Function to add message
def add_message(role: str, content: str):
    new_messages = messages.value + [{"role": role, "content": content}]
    messages.set(new_messages)

# Function to save the last assistant message to a docx file
def save_last_response_to_docx():
    doc = Document()
    doc.add_heading('Solara LLaMA Chat Last Response', level=1)

    last_message = next((msg for msg in reversed(messages.value) if msg["role"] == "assistant"), None)
    if last_message:
        doc.add_heading('LLaMA:', level=2)
        doc.add_paragraph(last_message["content"])
    
    file_path = "LLaMA_Last_Response.docx"
    doc.save(file_path)
    return file_path

# Function to call LLaMA API and get response
def call_llama():
    if len(messages.value) == 0 or messages.value[-1]["role"] != "user":
        return
    
    # Extract the user messages as a single input text
    user_messages = [msg["content"] for msg in messages.value if msg["role"] == "user"]
    input_text = "\n".join(user_messages)
    
    try:
        response = ollama_client.invoke(input=input_text)
        # print(response)
        assistant_message = response
        add_message("assistant", "")  # Add empty assistant message first
        threading.Thread(target=display_typing_effect, args=(assistant_message,)).start()
    except Exception as e:
        add_message("assistant", f"Error: {str(e)}")

# Function to display typing effect
def display_typing_effect(content: str):
    typing_message = ""
    for char in content:
        typing_message += char
        time.sleep(0.05)  # Adjust typing speed here
        messages.set(messages.value[:-1] + [{"role": "assistant", "content": typing_message}])
    typing_complete.set(True)  # Mark typing as complete
    save_last_response_to_docx()  # Save only the last assistant's response

@sol.component
def Home():
    # def on_button_click(topic):
    #     add_message("user", topic)
    #     call_llama()
    sol.Style(Path("style.css"))

    with sol.Column(
        style={"width": "90%", "height": "70vh"},
    ):
        # sol.AppBar(children=[sol.Text("AI Agent")])

        with sol.Row(justify="center"):
            with sol.Column():
                with sol.Link("/chat"):
                    sol.Button("Search and Create a .docx Wikipeia", outlined=True, classes=["btn"])
                    sol.Button("Photo organiser", outlined=True, classes=["btn"])

@sol.component
def Chat():
    sol.Style(Path("style.css"))

    user_message_count = len([m for m in messages.value if m["role"] == "user"])

    # Function to handle user input
    def send(message):
        add_message("user", message)
        call_llama()
        typing_complete.set(False)  # Reset typing complete state

    with sol.Column(
        style={"width": "90%", "height": "70vh"},
    ):
        # sol.AppBar(children=[sol.Text("AI Agent")])

        with sol.lab.ChatBox():
            for item in messages.value:
                with sol.lab.ChatMessage(
                    user=item["role"] == "user",
                    avatar=False,
                    name="LLaMA" if item["role"] == "assistant" else "User",
                    color="rgba(0,0,0, 0.06)" if item["role"] == "assistant" else "#ff991f",
                    avatar_background_color="primary" if item["role"] == "assistant" else None,
                    border_radius="20px",
                ):
                    sol.Markdown(item["content"])
        sol.lab.ChatInput(send_callback=send, disabled=False)

        if typing_complete.value and user_message_count > 0:
            file_path = save_last_response_to_docx()
            with open(file_path, "rb") as file:
                file_content = file.read()
            sol.FileDownload(file_content, label="Download Last Response", filename="LLaMA_Last_Response.docx")

routes = [
    sol.Route(path="/", component=Home, label="Home"),
    sol.Route(path="chat", component=Chat, label="Chat"),
]